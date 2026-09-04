import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    h.paragraph_format.keep_with_next = True
    if level == 1:
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(79, 70, 229) # Indigo
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(13, 148, 136) # Teal
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
    return h

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    
    # Border & background via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="4F46E5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:pBdr>')
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
    pPr.append(pBdr)
    pPr.append(shd)

    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 41, 59)

def add_callout(doc, text, title="NOTE"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="0D9488"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:pBdr>')
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF2FF"/>')
    pPr.append(pBdr)
    pPr.append(shd)

    r_title = p.add_run(f"[{title}] ")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(79, 70, 229)

    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)

def build_profile_doc(output_path, screenshot1_path, screenshot2_path):
    doc = docx.Document()

    # Set Margins to 0.75 in
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(30, 41, 59)

    # ==========================================================
    # COVER / HEADER BLOCK
    # ==========================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    run_super = p_title.add_run("DEPARTMENT OF COMPUTER SCIENCE & ARTIFICIAL INTELLIGENCE\n")
    run_super.font.name = 'Calibri'
    run_super.font.size = Pt(11)
    run_super.font.bold = True
    run_super.font.color.rgb = RGBColor(100, 116, 139)

    run_title = p_title.add_run("FLUTTER PROFILE CARD SCREEN IMPLEMENTATION\n")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(79, 70, 229)

    run_sub = p_title.add_run("Assignment 3: Responsive UI Layout, Core Widgets & Material 3 Theming")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(13, 148, 136)

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Student Name:", "R Virshin"),
        ("Roll Number / ID:", "150096724147"),
        ("Course / Subject:", "Cross-Platform Mobile Application Development (Flutter & Dart)"),
        ("Environment / Framework:", "Flutter 3.x • Dart 3.x • Material 3 • Chrome Web / macOS"),
    ]
    for row_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        r_l = p_lbl.add_run(label)
        r_l.font.bold = True
        r_l.font.color.rgb = RGBColor(79, 70, 229)
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(2)
        r_v = p_val.add_run(val)
        r_v.font.color.rgb = RGBColor(30, 41, 59)
        
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, top=60, bottom=60, left=120, right=120)
        set_cell_margins(cell_val, top=60, bottom=60, left=120, right=120)

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(8)
    p_div.paragraph_format.space_after = Pt(12)

    # ==========================================================
    # SECTION 1: EXECUTIVE SUMMARY & OBJECTIVES
    # ==========================================================
    add_styled_heading(doc, "1. Executive Summary & Assignment Objectives", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "This project showcases a production-ready, highly polished Profile Card screen built with the Flutter SDK. "
        "The objective of Assignment 3 is to demonstrate mastery over core Flutter layout paradigms, responsive container styling, "
        "idiomatic widget composition, and custom theme architecture. The implementation strictly adheres to the patterns established "
        "within the Cross-App curriculum, integrating clean architectural domain modeling and sound Dart 3 null safety."
    )

    p_obj = doc.add_paragraph()
    p_obj.add_run("Key technical objectives fulfilled in this assignment:\n")
    bullet_items = [
        ("Foundational Layout Widgets: ", "Seamlessly combining Column and Row widgets to establish primary and cross-axis alignment with zero overflow issues."),
        ("Card & Container Architecture: ", "Crafting an elegant, modern profile card using Container with border radius, subtle borders, and diffused drop shadows."),
        ("Profile Avatar Design: ", "Engineering a multi-tiered circular avatar utilizing concentric CircleAvatar widgets and framed borders."),
        ("Typographic Scale & Icons: ", "Displaying clear visual hierarchy with customized Text styles and contextual Material Icons."),
        ("Custom Theme Colors (AppPalette): ", "Establishing a dedicated color palette with semantic names and integrating it into Material 3 ThemeData."),
        ("Defensive Null Safety & Domain Entity: ", "Decoupling student profile information into an immutable UserProfile model with null-coalescing getters."),
        ("Automated Quality Assurance: ", "Ensuring 100% test pass rate in flutter test and 0 warnings or lints in flutter analyze.")
    ]
    for b_title, b_desc in bullet_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        r1 = bp.add_run(b_title)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(79, 70, 229)
        bp.add_run(b_desc)

    # ==========================================================
    # SECTION 2: ARCHITECTURE & WIDGET TREE BREAKDOWN
    # ==========================================================
    add_styled_heading(doc, "2. Architecture & Widget Tree Breakdown", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "The application follows a structured, clean architectural pattern separating design tokens (AppPalette), data entities (UserProfile), "
        "application configuration (ProfileApp), and user interface presentation (ProfileCardScreen with reusable helper components). "
        "The structural widget hierarchy is depicted below:"
    )

    add_code_block(doc, 
"""MaterialApp (theme: ThemeData with AppPalette.primary, AppPalette.secondary)
 └── Scaffold (backgroundColor: AppPalette.background)
      ├── AppBar (Title: 'Profile Card', Actions: [Share Icon])
      └── Center
           └── SingleChildScrollView (Padding: EdgeInsets.all(20))
                └── Column (Main Layout Wrapper)
                     └── Container (Main Profile Card: rounded corners, border, shadow)
                          └── Column (Card Content: MainAxisSize.min)
                               ├── Container (Concentric Avatar Frame)
                               │    └── CircleAvatar (Light Indigo background)
                               │         └── CircleAvatar (Primary Indigo + Person Icon)
                               ├── Text ('R Virshin' - Headline 22pt Bold)
                               ├── Text ('Flutter & Mobile App Developer' - Subtitle 14pt SemiBold)
                               ├── Row (Badges: 'Developer' & 'Verified' Container Pills)
                               ├── Container (Biography Box with rounded corners)
                               │    └── Text (profile.displayBio with line height 1.45)
                               ├── Container (Statistics Panel with dividers)
                               │    └── Row (Projects [12], Repos [35], Rating [4.9 ★])
                               ├── Column (Contact Details Tile List)
                               │    ├── _ContactInfoRow (Roll Number: 150096724147)
                               │    ├── _ContactInfoRow (Department: Computer Science & AI)
                               │    ├── _ContactInfoRow (Email: virshinkumar@gmail.com)
                               │    ├── _ContactInfoRow (Phone: +91 98765 43210)
                               │    └── _ContactInfoRow (Location: Bangalore, India)
                               └── Row (Interactive Action Buttons)
                                    ├── Expanded -> Container ('Message' Button with Send Icon)
                                    └── Expanded -> Container ('Connect' Button with Add Person Icon)""")

    # Required Widgets Table
    add_styled_heading(doc, "Core Required Widgets & Implementation Details", level=2)
    w_table = doc.add_table(rows=7, cols=3)
    w_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Widget", "Curriculum Role", "Implementation in Profile Card"]
    for idx, h_text in enumerate(headers):
        cell = w_table.rows[0].cells[idx]
        set_cell_background(cell, "4F46E5")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    widget_matrix = [
        ("Column", "Vertical layout structuring", "Organizes the avatar, name, subtitle, badges, bio, stats, contact list, and buttons into an aesthetic vertical sequence."),
        ("Row", "Horizontal layout structuring", "Distributes status badge pills, three-column statistics metrics, contact icon-label pairings, and dual action buttons."),
        ("Container", "Box styling & decoration", "Constructs the primary elevated card (BorderRadius.circular(24), BoxShadow, Border), badge chips, and action buttons."),
        ("CircleAvatar", "Circular clipping & avatars", "Layers two concentric avatars (radius 46 and 42) with contrast theme backgrounds to display the primary profile icon."),
        ("Text", "Typographic rendering", "Renders name (22pt bold), role (14pt semibold), bio (12.5pt muted), statistics, contact values, and button labels."),
        ("Icon", "Vector Material iconography", "Provides intuitive icons including person, code, verified badge, star, email, phone, location, and action glyphs.")
    ]

    for row_idx, (w_name, w_role, w_impl) in enumerate(widget_matrix, start=1):
        row = w_table.rows[row_idx]
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate([w_name, w_role, w_impl]):
            c = row.cells[col_idx]
            set_cell_background(c, bg)
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(79, 70, 229)

    doc.add_page_break()

    # ==========================================================
    # SECTION 3: THEME SYSTEM & DESIGN PALETTE
    # ==========================================================
    add_styled_heading(doc, "3. Custom Theme Architecture & Color Palette", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "A standout characteristic of professional Flutter applications is intentional, curated color selection. "
        "Rather than relying on default generic colors (e.g., Colors.blue or Colors.purple), this assignment implements a dedicated "
        "AppPalette class. This design pattern mirrors enterprise-grade Flutter architecture, enabling global updates from a single point of truth."
    )

    palette_table = doc.add_table(rows=11, cols=4)
    palette_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pal_headers = ["Constant Name", "Hex Code", "Preview / Color Role", "Semantic Purpose"]
    for idx, h_text in enumerate(pal_headers):
        cell = palette_table.rows[0].cells[idx]
        set_cell_background(cell, "0D9488")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    palette_data = [
        ("AppPalette.primary", "#4F46E5", "Indigo Accent", "Brand primary, AppBar background, primary avatar, and active action buttons"),
        ("AppPalette.primaryLight", "#EEF2FF", "Soft Indigo Tint", "Outer avatar ring, badge pill backgrounds, and contact icon frames"),
        ("AppPalette.secondary", "#0D9488", "Ocean Teal Accent", "Secondary brand accent, verification checkmarks, and secondary pill text"),
        ("AppPalette.secondaryLight", "#CCFBF1", "Soft Teal Tint", "Verification pill background creating high contrast readability"),
        ("AppPalette.background", "#F1F5F9", "Slate-100 Background", "Scaffold canvas color providing subtle contrast against the white card"),
        ("AppPalette.surface", "#FFFFFF", "Pure White Card", "Card body surface reflecting clean elevation above the canvas"),
        ("AppPalette.cardBorder", "#E2E8F0", "Slate-200 Border", "Subtle border outlining cards, stats panel, and contact tiles"),
        ("AppPalette.textPrimary", "#1E293B", "Slate-800 Heading", "High-contrast text for student name, primary labels, and contact values"),
        ("AppPalette.textSecondary", "#64748B", "Slate-500 Body", "Muted body text for biography, section descriptors, and contact titles"),
        ("AppPalette.star", "#F59E0B", "Amber Star", "Vibrant star icon color for user review rating display"),
    ]

    for row_idx, (p_name, p_hex, p_role, p_purp) in enumerate(palette_data, start=1):
        row = palette_table.rows[row_idx]
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate([p_name, p_hex, p_role, p_purp]):
            c = row.cells[col_idx]
            set_cell_background(c, bg)
            set_cell_margins(c, top=50, bottom=50, left=90, right=90)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            if col_idx == 0:
                r.font.bold = True
                r.font.name = 'Consolas'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(79, 70, 229)
            elif col_idx == 1:
                r.font.name = 'Consolas'
                r.font.size = Pt(8.5)

    add_callout(doc, 
        "Material 3 Integration: The AppPalette constants are directly wired into ThemeData using ColorScheme.fromSeed(seedColor: AppPalette.primary), "
        "allowing all standard Material 3 components to derive harmonious primary, secondary, and surface tonalities automatically.",
        title="DESIGN SYSTEM NOTE"
    )

    # ==========================================================
    # SECTION 4: DOMAIN MODELING & NULL SAFETY
    # ==========================================================
    add_styled_heading(doc, "4. Domain Modeling & Defensive Null Safety", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "To adhere strictly to Clean Code principles, student profile information is decoupled from UI layout code through an immutable "
        "UserProfile domain model. The class encapsulates sound null-safety patterns acquired in earlier Dart modules:"
    )

    add_code_block(doc,
"""class UserProfile {
  final String name;
  final String role;
  final String rollNumber;
  final String department;
  final String email;
  final String phone;
  final String location;
  final String? bio;           // Nullable string demonstration
  final int projectsCount;
  final int repoCount;
  final double rating;

  const UserProfile({
    required this.name,
    required this.role,
    required this.rollNumber,
    required this.department,
    required this.email,
    required this.phone,
    required this.location,
    this.bio,
    this.projectsCount = 0,
    this.repoCount = 0,
    this.rating = 5.0,
  });

  /// Null-coalescing fallback getter ensuring graceful default handling
  String get displayBio =>
      bio ?? 'Passionate student developer building modern cross-platform apps.';
}""")

    p_model = doc.add_paragraph()
    p_model.add_run(
        "By defining displayBio with the null-coalescing operator (??), the UI is guaranteed never to encounter runtime null pointer exceptions "
        "or render ungraceful blank sections when backend payloads omit biography strings."
    )

    doc.add_page_break()

    # ==========================================================
    # SECTION 5: VISUAL PROOF & SCREENSHOT ANALYSIS
    # ==========================================================
    add_styled_heading(doc, "5. Visual Proof & Screen Capture Analysis", level=1)

    p_vis = doc.add_paragraph()
    p_vis.add_run(
        "The application was executed in a live environment (Google Chrome web renderer on macOS) at localhost:59235. "
        "The resulting user interface demonstrates pixel-perfect alignment, rich color harmony, and robust scrolling responsiveness. "
        "Below are the verified execution screenshots captured from the running application:"
    )

    # Screenshot 1
    add_styled_heading(doc, "Figure 1: Upper Profile Card & Navigation View", level=2)
    p_img1 = doc.add_paragraph()
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.paragraph_format.space_before = Pt(4)
    p_img1.paragraph_format.space_after = Pt(2)
    if os.path.exists(screenshot1_path):
        p_img1.add_run().add_picture(screenshot1_path, width=Inches(5.6))
    else:
        p_img1.add_run("[Screenshot 1 Missing at path]")

    p_cap1 = doc.add_paragraph()
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap1.paragraph_format.space_after = Pt(8)
    r_cap1 = p_cap1.add_run("Screenshot 1: Live execution showing custom Indigo AppBar, Concentric Avatar, Name, Badges, Bio Box, and Statistics Grid.")
    r_cap1.font.italic = True
    r_cap1.font.size = Pt(8.5)
    r_cap1.font.color.rgb = RGBColor(100, 116, 139)

    p_desc1 = doc.add_paragraph()
    p_desc1.add_run(
        "Analysis of Figure 1:\n"
        "• AppBar: Displays centered 'Profile Card' title with white typography and a right-aligned share action icon.\n"
        "• Concentric Avatar: Features a 3px Indigo border, surrounded by a soft #EEF2FF tinted inner circle, containing a primary CircleAvatar with the person icon.\n"
        "• Typography & Role: 'R Virshin' is rendered prominently in 22pt bold text, followed by the role 'Flutter & Mobile App Developer' in primary Indigo.\n"
        "• Badge Row: Dual rounded pills display 'Developer' with a code icon and 'Verified' with a teal verification badge.\n"
        "• Biography & Statistics: The bio container provides padded text with line height 1.45. Below it, the statistics panel cleanly divides Projects (12), Repos (35), and Rating (4.9 with amber star)."
    )

    doc.add_page_break()

    # Screenshot 2
    add_styled_heading(doc, "Figure 2: Scrolled View with Complete Contact Details & Interactive Action Buttons", level=2)
    p_img2 = doc.add_paragraph()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.paragraph_format.space_before = Pt(4)
    p_img2.paragraph_format.space_after = Pt(2)
    if os.path.exists(screenshot2_path):
        p_img2.add_run().add_picture(screenshot2_path, width=Inches(5.6))
    else:
        p_img2.add_run("[Screenshot 2 Missing at path]")

    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap2.paragraph_format.space_after = Pt(8)
    r_cap2 = p_cap2.add_run("Screenshot 2: Scrolled view displaying the complete 5-item contact list (Roll No, Dept, Email, Phone, Location) and dual Action Buttons.")
    r_cap2.font.italic = True
    r_cap2.font.size = Pt(8.5)
    r_cap2.font.color.rgb = RGBColor(100, 116, 139)

    p_desc2 = doc.add_paragraph()
    p_desc2.add_run(
        "Analysis of Figure 2:\n"
        "• Complete Contact List: Five individual tiles built using the reusable _ContactInfoRow component. "
        "Each tile contains a themed icon inside an #EEF2FF square container, a subtitle label in Slate-500, and bold data values in Slate-800:\n"
        "   - Roll Number: 150096724147 (Badge icon)\n"
        "   - Department: Computer Science & AI (School graduation cap icon)\n"
        "   - Email: virshinkumar@gmail.com (Email envelope icon)\n"
        "   - Phone: +91 98765 43210 (Phone call icon)\n"
        "   - Location: Bangalore, India (Map location marker pin icon)\n"
        "• Dual Action Buttons: Arranged in a balanced horizontal Row with Expanded wrappers:\n"
        "   - 'Message' Button: Filled Indigo container with white text and send icon.\n"
        "   - 'Connect' Button: Outlined white container with 1.5px Indigo border, primary text, and person-add icon."
    )

    # ==========================================================
    # SECTION 6: VERIFICATION & TESTING RESULTS
    # ==========================================================
    add_styled_heading(doc, "6. Static Analysis & Widget Testing Suite", level=1)

    p_test = doc.add_paragraph()
    p_test.add_run(
        "To satisfy strict academic and industry engineering standards, the project underwent rigorous static analysis "
        "using the latest Flutter Lints ruleset (version 6.0.0) and automated widget testing."
    )

    add_styled_heading(doc, "Static Analysis (flutter analyze)", level=2)
    add_code_block(doc,
"""$ flutter analyze
Analyzing profile...                                            
No issues found! (ran in 4.2s)""")

    add_styled_heading(doc, "Automated Component Test Suite (flutter test)", level=2)
    add_code_block(doc,
"""$ flutter test
00:00 +0: loading /Users/virshin/VScode/Cross-App/Assignments/profile/test/widget_test.dart
00:00 +0: Profile card renders required widgets and content
00:02 +1: All tests passed!""")

    p_test_desc = doc.add_paragraph()
    p_test_desc.add_run(
        "The automated test suite in test/widget_test.dart asserts that every mandatory widget type (Column, Row, Container, "
        "CircleAvatar, Text, and Icon) exists in the widget tree and validates all critical string values (Student Name, Roll Number, "
        "Email, Phone, and Location) without rendering overflows."
    )

    # ==========================================================
    # SECTION 7: CONCLUSION
    # ==========================================================
    add_styled_heading(doc, "7. Conclusion & Learning Outcomes", level=1)

    p_conc = doc.add_paragraph()
    p_conc.add_run(
        "The Assignment 3 Profile Card project successfully demonstrates full competence in Flutter's foundational layout system. "
        "By structuring the user interface with cohesive widget hierarchies, applying custom Material 3 color palettes, and adhering "
        "to clean code conventions, the resulting application is both visually stunning and technically robust. "
        "All requirements—including custom theme colors, responsive Column and Row layouts, rounded Container styling, dual CircleAvatars, "
        "and integrated iconography—have been thoroughly implemented, tested, and visually confirmed."
    )

    doc.save(output_path)
    print(f"Document successfully created and saved to: {output_path}")

if __name__ == '__main__':
    base_dir = '/Users/virshin/VScode/Cross-App/Assignments/profile'
    out_docx = os.path.join(base_dir, 'DOCUMENTATION.docx')
    s1 = os.path.join(base_dir, 'assets', 'screenshot_top.png')
    s2 = os.path.join(base_dir, 'assets', 'screenshot_bottom.png')
    build_profile_doc(out_docx, s1, s2)
